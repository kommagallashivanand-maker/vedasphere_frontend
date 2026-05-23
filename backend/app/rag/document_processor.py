import base64
import logging
from pathlib import Path
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp"}

_IMAGE_PROMPT = (
    "You are analyzing a project-related image (e.g., architecture diagram, system design, "
    "flowchart, ERD, UI mockup, screenshot, chart, or any visual document). "
    "Provide a detailed description covering: all components/services/elements shown, "
    "their relationships and data flows, labels and annotations, technology stack references, "
    "and any key design decisions visible. "
    "Be thorough — this description will be used to answer questions about the project."
)


# ── Vision helpers ─────────────────────────────────────────────────────────────

def _describe_image_bytes(image_bytes: bytes, mime_type: str) -> str:
    """Send raw image bytes to the active vision LLM. Returns empty string on failure."""
    from app.config import settings
    try:
        image_data = base64.b64encode(image_bytes).decode("utf-8")
        if settings.AI_PROVIDER == "openai":
            from openai import OpenAI
            client = OpenAI(api_key=settings.OPENAI_API_KEY)
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": [
                    {"type": "text", "text": _IMAGE_PROMPT},
                    {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_data}"}},
                ]}],
                max_tokens=1500,
            )
            return response.choices[0].message.content or ""
        else:
            import requests
            url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"
            body = {"contents": [{"parts": [
                {"text": _IMAGE_PROMPT},
                {"inline_data": {"mime_type": mime_type, "data": image_data}},
            ]}]}
            resp = requests.post(url, json=body, params={"key": settings.GOOGLE_API_KEY}, timeout=60)
            resp.raise_for_status()
            return resp.json()["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        logger.warning(f"Vision API failed: {e}")
        return ""


# ── PDF extractor (text + embedded images per page) ────────────────────────────

def _extract_page_images(fitz_doc, page_num: int) -> List[str]:
    """
    Extract and describe all meaningful images embedded on a PDF page.
    Returns a list of description strings. Skips images under 5KB (icons/decorations).
    """
    page = fitz_doc[page_num]
    image_list = page.get_images(full=True)
    descriptions = []

    mime_map = {
        "png": "image/png", "jpeg": "image/jpeg", "jpg": "image/jpeg",
        "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp",
        "tiff": "image/tiff",
    }

    for idx, img_info in enumerate(image_list):
        xref = img_info[0]
        try:
            base_image = fitz_doc.extract_image(xref)
            img_bytes = base_image["image"]
            ext = base_image["ext"].lower()

            if len(img_bytes) < 5000:
                logger.debug(f"Page {page_num+1} image {idx+1}: skipped (too small, {len(img_bytes)} bytes)")
                continue

            mime_type = mime_map.get(ext, "image/png")
            logger.info(f"Page {page_num+1} image {idx+1}: describing ({len(img_bytes)} bytes, {ext})")

            desc = _describe_image_bytes(img_bytes, mime_type)
            if desc.strip():
                descriptions.append(f"[Image {idx+1} on page {page_num+1}]: {desc.strip()}")
                logger.info(f"Page {page_num+1} image {idx+1}: described ({len(desc)} chars)")
            else:
                logger.warning(f"Page {page_num+1} image {idx+1}: vision API returned empty")

        except Exception as e:
            logger.warning(f"Page {page_num+1} image {idx+1} (xref={xref}): failed — {e}")

    return descriptions


def extract_text_from_pdf(filepath: str) -> List[Document]:
    """
    Extract text AND embedded images from every PDF page.
    Each page produces one Document with:
      - The text layer
      - Vision LLM descriptions of any embedded images on that page
    """
    import fitz  # PyMuPDF
    from pypdf import PdfReader

    filename = Path(filepath).name
    pypdf_reader = PdfReader(filepath)
    fitz_doc = fitz.open(filepath)
    docs = []
    total_pages = len(pypdf_reader.pages)

    for page_num in range(total_pages):
        # Text layer
        text = (pypdf_reader.pages[page_num].extract_text() or "").strip()

        # Embedded images
        image_descriptions = _extract_page_images(fitz_doc, page_num)

        # Combine
        parts = []
        if text:
            parts.append(text)
        parts.extend(image_descriptions)

        if not parts:
            logger.debug(f"Page {page_num+1}: blank, skipping")
            continue

        combined = "\n\n".join(parts)
        docs.append(Document(
            page_content=combined,
            metadata={
                "source": filename,
                "page": page_num + 1,
                "has_images": len(image_descriptions) > 0,
                "image_count": len(image_descriptions),
            },
        ))
        logger.info(
            f"{filename} page {page_num+1}/{total_pages}: "
            f"{len(text)} text chars, {len(image_descriptions)} image(s)"
        )

    fitz_doc.close()
    return docs


# ── Other text extractors ──────────────────────────────────────────────────────

def extract_text_from_docx(filepath: str) -> List[Document]:
    from docx import Document as DocxDocument
    doc = DocxDocument(filepath)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(
        page_content=full_text,
        metadata={"source": Path(filepath).name, "page": 1},
    )]


def extract_text_from_txt(filepath: str) -> List[Document]:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [Document(
        page_content=text,
        metadata={"source": Path(filepath).name, "page": 1},
    )]


# ── Standalone image file extractor ───────────────────────────────────────────

def extract_text_from_image(filepath: str) -> List[Document]:
    """Describe a standalone image file using the active vision LLM."""
    filename = Path(filepath).name
    mime_map = {
        "jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
        "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp",
        "tiff": "image/tiff", "tif": "image/tiff",
    }
    ext = Path(filepath).suffix.lower().lstrip(".")
    mime_type = mime_map.get(ext, "image/png")

    with open(filepath, "rb") as f:
        image_bytes = f.read()

    logger.info(f"Describing standalone image: {filename}")
    description = _describe_image_bytes(image_bytes, mime_type)

    if not description.strip():
        raise ValueError(f"Vision API returned empty description for {filename}")

    return [Document(
        page_content=description,
        metadata={"source": filename, "page": 1, "content_type": "image"},
    )]


# ── Dispatcher ─────────────────────────────────────────────────────────────────

def load_document(filepath: str) -> List[Document]:
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(filepath)
    elif ext == ".txt":
        return extract_text_from_txt(filepath)
    elif ext in IMAGE_EXTENSIONS:
        return extract_text_from_image(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def split_documents(
    docs: List[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(docs)
